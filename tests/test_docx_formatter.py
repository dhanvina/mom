"""
Unit tests for the DocxFormatter class.
"""

import os
import pytest
import tempfile
from unittest.mock import patch, MagicMock
from docx import Document
from app.formatter.docx_formatter import DocxFormatter

class TestDocxFormatter:
    """Test cases for DocxFormatter class."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = DocxFormatter()
        
        # Create test MoM data
        self.mom_data = {
            "meeting_title": "Test Meeting",
            "date_time": "2023-01-01 10:00",
            "location": "Conference Room A",
            "attendees": ["John Doe", "Jane Smith", "Bob Johnson"],
            "agenda": ["Item 1", "Item 2", "Item 3"],
            "discussion_points": ["Discussion 1", "Discussion 2"],
            "action_items": [
                {"description": "Task 1", "assignee": "John Doe", "deadline": "2023-01-15", "status": "pending"},
                {"description": "Task 2", "assignee": "Jane Smith", "deadline": "2023-01-20", "status": "in_progress"}
            ],
            "decisions": [
                {"decision": "Decision 1", "context": "Context for decision 1"},
                {"decision": "Decision 2", "context": "Context for decision 2"}
            ],
            "next_steps": ["Step 1", "Step 2"]
        }
    
    def test_init(self):
        """Test initialization with default and custom config."""
        # Default config
        formatter = DocxFormatter()
        assert formatter.config == {}
        
        # Custom config
        custom_config = {
            "logo_path": "/path/to/logo.png",
            "company_name": "Example Corp",
            "include_branding": True
        }
        formatter = DocxFormatter(config=custom_config)
        assert formatter.config == custom_config
    
    def test_format_basic(self):
        """Test formatting basic MoM data into Word document."""
        # Format MoM data
        docx_bytes = self.formatter.format(self.mom_data)
        
        # Verify result
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        
        # Save to temporary file and verify content
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(docx_bytes)
            temp_path = temp_file.name
        
        try:
            # Open the document and verify content
            doc = Document(temp_path)
            
            # Check title
            assert doc.paragraphs[0].text == "Test Meeting"
            
            # Check for section headings
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
            assert "Attendees" in headings
            assert "Agenda" in headings
            assert "Discussion Points" in headings
            assert "Action Items" in headings
            assert "Decisions" in headings
            assert "Next Steps" in headings
            
            # Check for tables (action items)
            assert len(doc.tables) > 0
        finally:
            # Clean up
            os.unlink(temp_path)
    
    def test_format_with_sentiment(self):
        """Test formatting MoM data with sentiment analysis into Word document."""
        # Add sentiment data
        mom_data = self.mom_data.copy()
        mom_data["sentiment"] = {
            "overall": {"score": 0.75, "label": "Positive"},
            "participants": {
                "John Doe": {"score": 0.8, "label": "Positive"},
                "Jane Smith": {"score": 0.6, "label": "Neutral"},
                "Bob Johnson": {"score": 0.9, "label": "Positive"}
            }
        }
        
        # Format MoM data
        docx_bytes = self.formatter.format(mom_data)
        
        # Verify result
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        
        # Save to temporary file and verify content
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(docx_bytes)
            temp_path = temp_file.name
        
        try:
            # Open the document and verify content
            doc = Document(temp_path)
            
            # Check for sentiment heading
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
            assert "Meeting Sentiment" in headings
            
            # Check for sentiment tables
            assert len(doc.tables) > 1  # Action items table + sentiment table
        finally:
            # Clean up
            os.unlink(temp_path)
    
    def test_format_with_analytics(self):
        """Test formatting MoM data with analytics into Word document."""
        # Add analytics data
        mom_data = self.mom_data.copy()
        mom_data["analytics"] = {
            "speaking_time": [
                {"name": "John Doe", "time": 300, "percentage": 50.0},
                {"name": "Jane Smith", "time": 180, "percentage": 30.0},
                {"name": "Bob Johnson", "time": 120, "percentage": 20.0}
            ],
            "topics": [
                {"name": "Project Status", "frequency": 15},
                {"name": "Budget", "frequency": 10},
                {"name": "Timeline", "frequency": 8}
            ],
            "meeting_duration": "10 minutes",
            "efficiency_score": "85%"
        }
        
        # Format MoM data
        docx_bytes = self.formatter.format(mom_data)
        
        # Verify result
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        
        # Save to temporary file and verify content
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(docx_bytes)
            temp_path = temp_file.name
        
        try:
            # Open the document and verify content
            doc = Document(temp_path)
            
            # Check for analytics heading
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
            assert "Meeting Analytics" in headings
            assert "Speaking Time Distribution" in headings
            assert "Main Topics Discussed" in headings
            
            # Check for analytics tables
            assert len(doc.tables) > 1  # Action items table + analytics table
        finally:
            # Clean up
            os.unlink(temp_path)
    
    def test_format_minimal_data(self):
        """Test formatting minimal MoM data into Word document."""
        # Minimal MoM data
        minimal_data = {
            "meeting_title": "Minimal Meeting"
        }
        
        # Format MoM data
        docx_bytes = self.formatter.format(minimal_data)
        
        # Verify result
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        
        # Save to temporary file and verify content
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(docx_bytes)
            temp_path = temp_file.name
        
        try:
            # Open the document and verify content
            doc = Document(temp_path)
            
            # Check title
            assert doc.paragraphs[0].text == "Minimal Meeting"
            
            # Check that no section headings are present
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
            assert len(headings) == 1  # Only the title
            assert "Attendees" not in headings
            assert "Agenda" not in headings
            
            # Check that no tables are present
            assert len(doc.tables) == 0
        finally:
            # Clean up
            os.unlink(temp_path)
    
    def test_format_string_action_items(self):
        """Test formatting MoM data with string action items into Word document."""
        # MoM data with string action items
        mom_data = self.mom_data.copy()
        mom_data["action_items"] = ["Do task 1", "Do task 2"]
        
        # Format MoM data
        docx_bytes = self.formatter.format(mom_data)
        
        # Verify result
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        
        # Save to temporary file and verify content
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(docx_bytes)
            temp_path = temp_file.name
        
        try:
            # Open the document and verify content
            doc = Document(temp_path)
            
            # Check for action items heading
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
            assert "Action Items" in headings
            
            # Check that no action items table is present (should be bullet points instead)
            # This is a bit tricky to test directly, so we'll check that there are fewer tables
            assert len(doc.tables) == 0
        finally:
            # Clean up
            os.unlink(temp_path)
    
    def test_format_string_decisions(self):
        """Test formatting MoM data with string decisions into Word document."""
        # MoM data with string decisions
        mom_data = self.mom_data.copy()
        mom_data["decisions"] = ["Decision 1", "Decision 2"]
        
        # Format MoM data
        docx_bytes = self.formatter.format(mom_data)
        
        # Verify result
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
    
    def test_format_with_branding(self):
        """Test formatting MoM data with branding into Word document."""
        # Create a temporary logo file
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_logo:
            temp_logo_path = temp_logo.name
        
        try:
            # Create formatter with branding config
            branding_config = {
                "logo_path": temp_logo_path,
                "company_name": "Example Corp",
                "include_branding": True
            }
            formatter = DocxFormatter(config=branding_config)
            
            # Format MoM data
            docx_bytes = formatter.format(self.mom_data)
            
            # Verify result
            assert isinstance(docx_bytes, bytes)
            assert len(docx_bytes) > 0
        finally:
            # Clean up
            os.unlink(temp_logo_path)
    
    def test_format_without_branding(self):
        """Test formatting MoM data without branding into Word document."""
        # Create formatter with branding disabled
        config = {"include_branding": False}
        formatter = DocxFormatter(config=config)
        
        # Format MoM data
        docx_bytes = formatter.format(self.mom_data)
        
        # Verify result
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
    
    def test_format_error(self):
        """Test error handling during formatting."""
        # Mock Document to raise an exception
        with patch('app.formatter.docx_formatter.Document', side_effect=Exception("Test error")):
            # Format MoM data
            with pytest.raises(ValueError) as excinfo:
                self.formatter.format(self.mom_data)
            
            # Verify error message
            assert "Error formatting MoM data into Word document" in str(excinfo.value)