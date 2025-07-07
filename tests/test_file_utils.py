import pytest
from unittest.mock import MagicMock, patch
from app.utils.file_utils import FileHandler
import io
import docx

class MockUploadedFile:
    def __init__(self, name, file_type, content):
        self.name = name
        self.type = file_type
        self._content = content
        self.size = len(content)

    def getvalue(self):
        return self._content

    def read(self):
        return io.BytesIO(self._content)

# ---------- Tests for read_file_content ----------

def test_read_txt_file_content():
    handler = FileHandler()
    file_content = b"Hello from text file."
    uploaded_file = MockUploadedFile("test.txt", "text/plain", file_content)

    result = handler.read_file_content(uploaded_file)
    assert result == "Hello from text file."

def test_read_docx_file_content(tmp_path):
    # Create a .docx file in memory
    doc = docx.Document()
    doc.add_paragraph("First line")
    doc.add_paragraph("Second line")
    fake_io = io.BytesIO()
    doc.save(fake_io)
    fake_io.seek(0)

    uploaded_file = MockUploadedFile(
        "test.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        fake_io.getvalue()
    )

    handler = FileHandler()
    result = handler.read_file_content(uploaded_file)
    assert "First line" in result
    assert "Second line" in result

def test_read_file_content_invalid_type():
    handler = FileHandler()
    uploaded_file = MockUploadedFile("test.pdf", "application/pdf", b"Some content")

    result = handler.read_file_content(uploaded_file)
    assert result == ""

def test_read_file_content_unicode_error():
    handler = FileHandler()
    # Invalid UTF-8 sequence to force UnicodeDecodeError
    uploaded_file = MockUploadedFile("bad.txt", "text/plain", b"\xff\xfe")

    result = handler.read_file_content(uploaded_file)
    assert result == ""

# ---------- Tests for get_file_info ----------

def test_get_file_info_valid_file():
    handler = FileHandler()
    uploaded_file = MockUploadedFile("file.txt", "text/plain", b"12345678")

    info = handler.get_file_info(uploaded_file)
    assert info["name"] == "file.txt"
    assert info["type"] == "text/plain"
    assert info["size"] == 8
    assert isinstance(info["size_mb"], float)

def test_get_file_info_none():
    handler = FileHandler()
    info = handler.get_file_info(None)
    assert info == {}

# ---------- Mocking Streamlit Uploader ----------

@patch("app.utils.file_utils.st.file_uploader")

def test_upload_file_returns_none(mock_uploader):
    handler = FileHandler()
    mock_uploader.return_value = None
    result = handler.upload_file()
    assert result is None

@patch("app.utils.file_utils.st.file_uploader")

def test_upload_file_success(mock_uploader):
    handler = FileHandler()
    uploaded_file = MockUploadedFile("demo.txt", "text/plain", b"test")
    mock_uploader.return_value = uploaded_file

    with patch.object(handler, "read_file_content", return_value="test") as mock_read:
        result = handler.upload_file()
        mock_read.assert_called_once_with(uploaded_file)
        assert result == "test"

# ---------- Mocking Streamlit Text Area ----------

@patch("app.utils.file_utils.st.text_area")

def test_text_input_area(mock_text_area):
    handler = FileHandler()
    mock_text_area.return_value = "Some transcript text"
    result = handler.text_input_area()
    assert result == "Some transcript text"
