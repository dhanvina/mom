"""
Docx formatter module for AI-powered MoM generator.

This module provides functionality for formatting MoM data into Word document format.
"""

import os
import logging
import tempfile
from typing import Dict, Any, Optional, List, Union, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .base_formatter import BaseFormatter

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocxFormatter(BaseFormatter):
    """
    Formats MoM data into Word document format.
    
    This class provides methods for formatting MoM data into Word documents
    using python-docx.
    
    Attributes:
        config (Dict): Configuration options for the formatter
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize the DocxFormatter with optional configuration.
        
        Args:
            config (Dict, optional): Configuration options for the formatter
        """
        super().__init__(config)
        logger.info("DocxFormatter initialized")
    
    def format(self, mom_data: Dict[str, Any]) -> bytes:
        """
        Format MoM data into a Word document.
        
        Args:
            mom_data (Dict[str, Any]): MoM data to format
            
        Returns:
            bytes: Word document as bytes
        """
        try:
            logger.info("Formatting MoM data into Word document")
            
            # Create a new document
            doc = Document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add title
            if "meeting_title" in mom_data:
                doc.add_heading(mom_data["meeting_title"], level=0)
            
            # Add metadata section
            self._add_metadata_section(doc, mom_data)
            
            # Add attendees section
            self._add_attendees_section(doc, mom_data)
            
            # Add agenda section
            self._add_agenda_section(doc, mom_data)
            
            # Add discussion points section
            self._add_discussion_section(doc, mom_data)
            
            # Add action items section
            self._add_action_items_section(doc, mom_data)
            
            # Add decisions section
            self._add_decisions_section(doc, mom_data)
            
            # Add next steps section
            self._add_next_steps_section(doc, mom_data)
            
            # Add analytics section if available
            if "analytics" in mom_data:
                self._add_analytics_section(doc, mom_data["analytics"])
            
            # Add sentiment section if available
            if "sentiment" in mom_data:
                self._add_sentiment_section(doc, mom_data["sentiment"])
            
            # Add branding if configured
            if self.config.get("include_branding", True):
                self._add_branding(doc, mom_data)
            
            # Save the document to a temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                docx_path = temp_file.name
            
            try:
                # Save the document
                doc.save(docx_path)
                
                # Read the document as bytes
                with open(docx_path, 'rb') as f:
                    docx_bytes = f.read()
                
                return docx_bytes
            
            finally:
                # Clean up temporary file
                if os.path.exists(docx_path):
                    os.unlink(docx_path)
        
        except Exception as e:
            error_msg = f"Error formatting MoM data into Word document: {e}"
            logger.error(error_msg)
            raise ValueError(error_msg)
    
    def _setup_document_styles(self, doc: Document) -> None:
        """
        Set up document styles.
        
        Args:
            doc (Document): Word document
        """
        # Set up styles
        styles = doc.styles
        
        # Modify heading styles
        heading_style = styles['Heading 1']
        heading_style.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        heading_style.font.size = Pt(16)
        
        heading_style = styles['Heading 2']
        heading_style.font.color.rgb = RGBColor(0x34, 0x98, 0xdb)
        heading_style.font.size = Pt(14)
        
        # Create custom styles
        if 'Action Item' not in styles:
            style = styles.add_style('Action Item', WD_STYLE_TYPE.PARAGRAPH)
            style.font.color.rgb = RGBColor(0x27, 0xae, 0x60)
            style.font.bold = True
        
        if 'Decision' not in styles:
            style = styles.add_style('Decision', WD_STYLE_TYPE.PARAGRAPH)
            style.font.color.rgb = RGBColor(0x8e, 0x44, 0xad)
            style.font.bold = True
        
        if 'Next Step' not in styles:
            style = styles.add_style('Next Step', WD_STYLE_TYPE.PARAGRAPH)
            style.font.color.rgb = RGBColor(0xe7, 0x4c, 0x3c)
            style.font.bold = True
    
    def _add_metadata_section(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add metadata section to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        if "date_time" in mom_data or "location" in mom_data:
            # Add a paragraph for metadata
            paragraph = doc.add_paragraph()
            
            if "date_time" in mom_data:
                run = paragraph.add_run("Date & Time: ")
                run.bold = True
                paragraph.add_run(mom_data["date_time"])
                paragraph.add_run("\n")
            
            if "location" in mom_data:
                run = paragraph.add_run("Location: ")
                run.bold = True
                paragraph.add_run(mom_data["location"])
            
            # Add spacing after metadata
            doc.add_paragraph()
    
    def _add_attendees_section(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add attendees section to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        if "attendees" not in mom_data or not mom_data["attendees"]:
            return
        
        # Add heading
        doc.add_heading("Attendees", level=1)
        
        # Check if attendees is a list of strings or a list of dictionaries
        if isinstance(mom_data["attendees"][0], str):
            # List of strings
            for attendee in mom_data["attendees"]:
                paragraph = doc.add_paragraph(style='List Bullet')
                paragraph.add_run(attendee)
        else:
            # List of dictionaries
            for attendee in mom_data["attendees"]:
                name = attendee.get("name", "")
                role = attendee.get("role", "")
                
                paragraph = doc.add_paragraph(style='List Bullet')
                if role:
                    paragraph.add_run(f"{name} ({role})")
                else:
                    paragraph.add_run(name)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_agenda_section(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add agenda section to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        if "agenda" not in mom_data or not mom_data["agenda"]:
            return
        
        # Add heading
        doc.add_heading("Agenda", level=1)
        
        # Add agenda items
        for item in mom_data["agenda"]:
            paragraph = doc.add_paragraph(style='List Bullet')
            paragraph.add_run(item)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_discussion_section(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add discussion points section to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        if "discussion_points" not in mom_data or not mom_data["discussion_points"]:
            return
        
        # Add heading
        doc.add_heading("Discussion Points", level=1)
        
        # Add discussion points
        for item in mom_data["discussion_points"]:
            paragraph = doc.add_paragraph(style='List Bullet')
            paragraph.add_run(item)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_action_items_section(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add action items section to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        if "action_items" not in mom_data or not mom_data["action_items"]:
            return
        
        # Add heading
        doc.add_heading("Action Items", level=1)
        
        # Check if action_items is a list of strings or a list of dictionaries
        if isinstance(mom_data["action_items"][0], str):
            # List of strings
            for item in mom_data["action_items"]:
                paragraph = doc.add_paragraph(style='List Bullet')
                run = paragraph.add_run(item)
                run.style = 'Action Item'
        else:
            # List of dictionaries - create a table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Action Item"
            header_cells[1].text = "Assignee"
            header_cells[2].text = "Due Date"
            header_cells[3].text = "Status"
            
            # Make header row bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add data rows
            for item in mom_data["action_items"]:
                description = item.get("description", item.get("task", ""))
                assignee = item.get("assignee", "")
                deadline = item.get("deadline", item.get("due", ""))
                status = item.get("status", "")
                
                row_cells = table.add_row().cells
                row_cells[0].text = description
                row_cells[1].text = assignee
                row_cells[2].text = deadline
                row_cells[3].text = status
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_decisions_section(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add decisions section to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        if "decisions" not in mom_data or not mom_data["decisions"]:
            return
        
        # Add heading
        doc.add_heading("Decisions", level=1)
        
        # Check if decisions is a list of strings or a list of dictionaries
        if isinstance(mom_data["decisions"][0], str):
            # List of strings
            for item in mom_data["decisions"]:
                paragraph = doc.add_paragraph(style='List Bullet')
                run = paragraph.add_run(item)
                run.style = 'Decision'
        else:
            # List of dictionaries
            for item in mom_data["decisions"]:
                decision = item.get("decision", "")
                context = item.get("context", "")
                
                paragraph = doc.add_paragraph(style='List Bullet')
                run = paragraph.add_run(decision)
                run.style = 'Decision'
                
                if context:
                    # Add context as a sublist item
                    paragraph = doc.add_paragraph(style='List Bullet 2')
                    run = paragraph.add_run("Context: ")
                    run.bold = True
                    paragraph.add_run(context)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_next_steps_section(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add next steps section to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        if "next_steps" not in mom_data or not mom_data["next_steps"]:
            return
        
        # Add heading
        doc.add_heading("Next Steps", level=1)
        
        # Add next steps
        for item in mom_data["next_steps"]:
            paragraph = doc.add_paragraph(style='List Bullet')
            run = paragraph.add_run(item)
            run.style = 'Next Step'
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_analytics_section(self, doc: Document, analytics: Dict[str, Any]) -> None:
        """
        Add analytics section to the document.
        
        Args:
            doc (Document): Word document
            analytics (Dict[str, Any]): Analytics data
        """
        if not analytics:
            return
        
        # Add heading
        doc.add_heading("Meeting Analytics", level=1)
        
        # Add speaking time distribution if available
        if "speaking_time" in analytics:
            doc.add_heading("Speaking Time Distribution", level=2)
            
            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Participant"
            header_cells[1].text = "Speaking Time"
            header_cells[2].text = "Percentage"
            
            # Make header row bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add data rows
            for item in analytics["speaking_time"]:
                name = item.get("name", "")
                time = item.get("time", 0)
                percentage = item.get("percentage", 0)
                
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = f"{time} seconds"
                row_cells[2].text = f"{percentage:.1f}%"
            
            # Add spacing after table
            doc.add_paragraph()
        
        # Add topics if available
        if "topics" in analytics:
            doc.add_heading("Main Topics Discussed", level=2)
            
            for topic in analytics["topics"]:
                if isinstance(topic, dict):
                    name = topic.get("name", "")
                    frequency = topic.get("frequency", 0)
                    
                    paragraph = doc.add_paragraph(style='List Bullet')
                    paragraph.add_run(f"{name} (mentioned {frequency} times)")
                else:
                    paragraph = doc.add_paragraph(style='List Bullet')
                    paragraph.add_run(topic)
            
            # Add spacing after topics
            doc.add_paragraph()
        
        # Add any other analytics as simple paragraphs
        for key, value in analytics.items():
            if key not in ["speaking_time", "topics"] and isinstance(value, str):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{key.replace('_', ' ').title()}: ")
                run.bold = True
                paragraph.add_run(value)
    
    def _add_sentiment_section(self, doc: Document, sentiment: Union[Dict[str, Any], str]) -> None:
        """
        Add sentiment section to the document.
        
        Args:
            doc (Document): Word document
            sentiment (Union[Dict[str, Any], str]): Sentiment data
        """
        if not sentiment:
            return
        
        # Add heading
        doc.add_heading("Meeting Sentiment", level=1)
        
        if isinstance(sentiment, dict):
            # Format sentiment data
            overall = sentiment.get("overall", {})
            overall_score = overall.get("score", 0)
            overall_label = overall.get("label", "Neutral")
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("Overall Sentiment: ")
            run.bold = True
            paragraph.add_run(f"{overall_label} ({overall_score:.2f})")
            
            # Add participant sentiments if available
            if "participants" in sentiment:
                doc.add_heading("Participant Sentiments", level=2)
                
                # Create table
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Participant"
                header_cells[1].text = "Sentiment"
                header_cells[2].text = "Score"
                
                # Make header row bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for name, values in sentiment["participants"].items():
                    label = values.get("label", "Neutral")
                    score = values.get("score", 0)
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = name
                    row_cells[1].text = label
                    row_cells[2].text = f"{score:.2f}"
        else:
            # Simple string sentiment
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("Overall Sentiment: ")
            run.bold = True
            paragraph.add_run(sentiment)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_branding(self, doc: Document, mom_data: Dict[str, Any]) -> None:
        """
        Add branding elements to the document.
        
        Args:
            doc (Document): Word document
            mom_data (Dict[str, Any]): MoM data
        """
        # Get branding from config
        logo_path = self.config.get("logo_path", "")
        company_name = self.config.get("company_name", "")
        
        # Add logo if available
        if logo_path and os.path.exists(logo_path):
            try:
                # Add logo to header
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
            except Exception as e:
                logger.warning(f"Failed to add logo to document: {e}")
        
        # Add company name to footer if available
        if company_name:
            try:
                # Add company name to footer
                section = doc.sections[0]
                footer = section.footer
                paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(f"Generated by AI-powered MoM Generator for {company_name}")
            except Exception as e:
                logger.warning(f"Failed to add company name to document: {e}")
    
    def _create_element(self, name: str) -> OxmlElement:
        """
        Create an OxmlElement with the given name.
        
        Args:
            name (str): Element name
            
        Returns:
            OxmlElement: Created element
        """
        return OxmlElement(name)
    
    def _create_attribute(self, element: OxmlElement, name: str, value: str) -> None:
        """
        Create an attribute on the given element.
        
        Args:
            element (OxmlElement): Element to add attribute to
            name (str): Attribute name
            value (str): Attribute value
        """
        element.set(qn(name), value)