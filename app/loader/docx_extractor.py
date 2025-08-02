"""
DOCX extractor module for AI-powered MoM generator.

This module provides functionality for extracting text from Word documents.
"""

import logging
import docx
import io
from typing import Any, Dict, Optional, Union
from .base_extractor import BaseExtractor

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocxExtractor(BaseExtractor):
    """
    Extracts text from Word documents.
    
    This class provides methods for loading and extracting text from .docx files.
    
    Attributes:
        config (Dict): Configuration options for the extractor
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize the DocxExtractor with optional configuration.
        
        Args:
            config (Dict, optional): Configuration options for the extractor
        """
        super().__init__(config)
        logger.info("DocxExtractor initialized")
    
    def extract(self, source: Union[str, bytes, io.BytesIO]) -> str:
        """
        Extract text from a Word document.
        
        Args:
            source (Union[str, bytes, io.BytesIO]): Path to the Word document,
                bytes content, or BytesIO object
            
        Returns:
            str: The extracted text
            
        Raises:
            IOError: If the document cannot be read
        """
        try:
            logger.info(f"Extracting text from Word document")
            
            # Handle different source types
            if isinstance(source, str):
                # Source is a file path
                doc = docx.Document(source)
            elif isinstance(source, bytes):
                # Source is bytes
                doc = docx.Document(io.BytesIO(source))
            elif isinstance(source, io.BytesIO):
                # Source is BytesIO
                doc = docx.Document(source)
            else:
                raise ValueError(f"Unsupported source type: {type(source)}")
            
            # Extract text from paragraphs
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += "\n" + cell.text
            
            return self.preprocess(content)
        
        except Exception as e:
            error_msg = f"Error extracting text from Word document: {e}"
            logger.error(error_msg)
            raise IOError(error_msg)